from __future__ import annotations

import csv
import io
import json
import math
import os
import re
from dotenv import load_dotenv
from dataclasses import dataclass
from typing import Any, Dict, List, Tuple, Union

import numpy as np
from sentence_transformers import SentenceTransformer

from llm.client import LLMClient
from llm.prompts import (
    BEHAVIORAL_CHAT_SYSTEM_PROMPT,
    CONVERSATION_AND_QUESTION_TYPE_INSTRUCTION,
    PARCEL_DATASET_SYSTEM_PROMPT,
    TABLE_DATASET_SYSTEM_PROMPT,
    get_personalisation_system_snippet,
)

load_dotenv()

# Groq on_demand tier has a small token-per-request limit (~6k). Keep prompts compact.
MAX_ROWS_PER_REQUEST = 12
MAX_ROW_TEXT_CHARS = 280
MAX_HISTORY_TURNS = 6
MAX_HISTORY_CHARS = 240

# Local embedding model (no API key). Only GROQ_API_KEY is used (for LLM answers).
RAG_EMBEDDING_MODEL = "all-MiniLM-L6-v2"


@dataclass
class RAGChunk:
    """
    Represents a single chunk (row) from the dataset with its embedding.
    """

    row_index: int
    text: str
    metadata: Dict[str, Any]
    embedding: List[float]


def _chunk_text(text: str, chunk_size: int = 1200, overlap: int = 180) -> List[str]:
    """
    Split long free-form text into overlapping chunks for better document RAG.
    """
    s = (text or "").strip()
    if not s:
        return []
    if len(s) <= chunk_size:
        return [s]

    chunks: List[str] = []
    start = 0
    n = len(s)
    while start < n:
        end = min(start + chunk_size, n)
        window = s[start:end]
        # Try to break on sentence/newline near the end to keep chunks readable.
        if end < n:
            split_at = max(window.rfind("\n\n"), window.rfind(". "), window.rfind("\n"))
            if split_at > int(chunk_size * 0.55):
                end = start + split_at + 1
                window = s[start:end]
        chunk = window.strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(0, end - overlap)
    return chunks


def parse_table_input(table: Union[str, List[Any]]) -> List[Dict[str, Any]]:
    """
    Parse user-provided table into a list of record dicts (one per row).
    Supports: JSON array of objects, CSV/TSV string (first row = headers), list of dicts.
    """
    if isinstance(table, list):
        if not table:
            return []
        first = table[0]
        if isinstance(first, dict):
            return [dict(r) for r in table if isinstance(r, dict)]
        if isinstance(first, (list, tuple)):
            # Rows as lists; first row = headers
            rows = [list(r) for r in table]
            headers = [str(h) for h in rows[0]]
            return [
                dict(zip(headers, row))
                for row in rows[1:]
                if len(row) >= len(headers)
            ]
        return []

    if isinstance(table, str):
        s = table.strip().lstrip("\ufeff")  # strip BOM from Excel export
        if not s:
            return []
        # Try JSON first (array of objects or array of arrays)
        if (s.startswith("[") and s.endswith("]")) or s.startswith("{"):
            try:
                data = json.loads(s)
                if isinstance(data, list):
                    return parse_table_input(data)
                return []
            except json.JSONDecodeError:
                pass
        # CSV or TSV
        delimiter = "\t" if "\t" in s.split("\n")[0] else ","
        reader = csv.reader(io.StringIO(s), delimiter=delimiter)
        rows = list(reader)
        if not rows:
            return []
        headers = [str(h).strip() or f"col_{i}" for i, h in enumerate(rows[0])]
        out = []
        for row in rows[1:]:
            vals = [str(c).strip() for c in row]
            if len(vals) < len(headers):
                vals += [""] * (len(headers) - len(vals))
            out.append(dict(zip(headers, vals[: len(headers)])))
        return out

    return []


class ExcelRAG:
    """
    In-memory RAG index built from a table (file or in-memory records).

    - From file: loads JSON/CSV at dataset_path (e.g. us_parcel_dataset_2000.json).
    - From memory: build_index() uses provided records (e.g. user-uploaded table).
    - Turns each row into textual representation; stores embeddings + metadata for retrieval.
    """

    # def makeToData(self, 
    # dataset_path: str| None = None
    # ) -> None:
    # if

    def __init__(
        self,
        dataset_path: str | None = None,
        records: List[Dict[str, Any]] | None = None,
    ) -> None:
        if records is None:
            if dataset_path is None or not os.path.exists(dataset_path):
                raise RuntimeError(
                    f"Dataset file not found at path: {dataset_path}" if dataset_path
                    else "Provide either dataset_path (existing file) or records (list of dicts)."
                )

        self._dataset_path = dataset_path
        self._records = records  # in-memory table when provided
        self._from_user_table = records is not None
        self._embedder = SentenceTransformer(RAG_EMBEDDING_MODEL)
        self._chunks: List[RAGChunk] = []
        self._schema_summary: str = ""

    # Columns to always include in row text even when empty, so the LLM can tell "no value" from "missing".
    # E.g. for "give me all the emails" the model can list only rows where Email is present.
    _KEY_COLUMNS_ALWAYS_SHOW = {"Email", "Project. Name", "Account Name", "Contact Name", "Stage", "Cell"}

    @staticmethod
    def _record_to_text(record: Dict[str, Any]) -> str:
        """
        Convert a record (dict) to a compact text description.
        Flattens "owners" junction list when present. For key columns (e.g. Email), always
        include them even when empty as "(not in data)" so the LLM can filter (e.g. list only rows with emails).
        """
        parts = []
        key_columns = ExcelRAG._KEY_COLUMNS_ALWAYS_SHOW
        for key, value in record.items():
            is_empty = value is None or value == ""
            if key == "owners" and isinstance(value, list):
                owner_parts = []
                for o in value:
                    if isinstance(o, dict):
                        name = o.get("owner_name") or str(o.get("owner_id", ""))
                        pct = o.get("ownership_percent")
                        pct_str = f" ({pct}%)" if pct is not None else ""
                        email = o.get("email") or ""
                        phone = o.get("phone") or ""
                        contact = []
                        if email:
                            contact.append(f"email {email}")
                        if phone:
                            contact.append(f"phone {phone}")
                        contact_str = " ".join(contact)
                        owner_parts.append(f"{name}{pct_str}" + (f" {contact_str}" if contact_str else ""))
                if owner_parts:
                    parts.append(f"owners: {', '.join(owner_parts)}")
            elif key in key_columns:
                if is_empty:
                    parts.append(f"{key}: (not in data)")
                else:
                    parts.append(f"{key}: {value}")
            elif not is_empty:
                parts.append(f"{key}: {value}")
        return " | ".join(parts)

    def build_index(self, max_rows: int | None = None) -> None:
        """
        Build the in-memory index from file (JSON/CSV) or from in-memory records.
        Records can be any list of dicts (e.g. parcel rows or user-provided table).
        """
        if self._records is not None:
            rows = self._records[:max_rows] if max_rows is not None else self._records
        else:
            path_lower = (self._dataset_path or "").lower()
            if path_lower.endswith(".csv") or path_lower.endswith(".txt"):
                try:
                    with open(self._dataset_path, encoding="utf-8") as f:
                        text = f.read()
                except UnicodeDecodeError:
                    with open(self._dataset_path, encoding="utf-8-sig") as f:
                        text = f.read()
                text = text.lstrip("\ufeff")
                rows = parse_table_input(text)
                if not rows:
                    raise RuntimeError("Dataset CSV/TSV produced no records.")
            else:
                with open(self._dataset_path, encoding="utf-8") as f:
                    data = json.load(f)
                if not isinstance(data, list):
                    raise RuntimeError("Dataset JSON must be an array of records.")
                rows = data
            rows = rows[:max_rows] if max_rows is not None else rows

        self._total_row_count = len(rows)
        self._schema_summary = self._infer_schema_summary(rows)
        chunks_list: List[RAGChunk] = []

        for idx, record in enumerate(rows):
            if not isinstance(record, dict):
                continue
            text = self._record_to_text(record)
            if not text:
                continue

            emb = self._embed(text)
            chunks_list.append(
                RAGChunk(
                    row_index=idx,
                    text=text,
                    metadata=dict(record),
                    embedding=emb,
                )
            )

        self._chunks = chunks_list

    @staticmethod
    def _infer_schema_summary(rows: List[Dict[str, Any]]) -> str:
        """
        Build a compact schema + relation-hints summary from the current table.
        This helps the LLM answer questions like "what are the relationships?"
        without seeing the full dataset each time.
        """
        if not rows:
            return "Schema summary unavailable (no rows)."

        sample_rows = rows[:500]
        cols_ordered: List[str] = []
        seen_cols: set[str] = set()
        non_empty_count: Dict[str, int] = {}
        unique_values: Dict[str, set[str]] = {}
        nested_list_keys: set[str] = set()

        for r in sample_rows:
            if not isinstance(r, dict):
                continue
            for k, v in r.items():
                col = str(k)
                if col not in seen_cols:
                    seen_cols.add(col)
                    cols_ordered.append(col)
                if v is None or v == "":
                    continue
                non_empty_count[col] = non_empty_count.get(col, 0) + 1
                if isinstance(v, list) and v and isinstance(v[0], dict):
                    nested_list_keys.add(col)
                    continue
                unique_values.setdefault(col, set()).add(str(v))

        def _base_name(col_name: str) -> str:
            c = col_name.strip().lower()
            c = c.replace(".", " ").replace("-", " ")
            if c.endswith("_id"):
                c = c[:-3]
            if c.endswith(" id"):
                c = c[:-3]
            return c.strip()

        lower_cols = {c.lower(): c for c in cols_ordered}
        key_candidates: List[str] = []
        for c in cols_ordered:
            filled = non_empty_count.get(c, 0)
            uniq = len(unique_values.get(c, set()))
            if filled > 0 and uniq == filled and (c.lower() == "id" or "id" in c.lower()):
                key_candidates.append(c)

        relation_hints: List[str] = []
        for nk in sorted(nested_list_keys):
            relation_hints.append(f"one-to-many from main record to nested '{nk}' entries")

        for c in cols_ordered:
            cl = c.lower()
            if cl == "id" or cl.endswith("_id") or cl.endswith(" id"):
                base = _base_name(c)
                if not base:
                    continue
                for cand in (
                    f"{base}_name",
                    f"{base} name",
                    f"{base}. name",
                    f"{base}",
                ):
                    if cand in lower_cols and lower_cols[cand] != c:
                        relation_hints.append(
                            f"'{c}' likely links to entity described by '{lower_cols[cand]}'"
                        )
                        break

        cols_preview = ", ".join(cols_ordered[:18])
        if len(cols_ordered) > 18:
            cols_preview += ", ..."
        keys_preview = ", ".join(key_candidates[:6]) if key_candidates else "none"
        rel_preview = "; ".join(relation_hints[:6]) if relation_hints else "no strong explicit key links detected"
        return (
            f"Columns ({len(cols_ordered)}): {cols_preview}. "
            f"Key candidates: {keys_preview}. "
            f"Relation hints: {rel_preview}."
        )

    def _embed(self, text: str) -> List[float]:
        """Get an embedding vector using the local sentence-transformers model."""
        vec = self._embedder.encode(text, convert_to_numpy=True)
        return vec.tolist()

    @staticmethod
    def _to_float(value: Any) -> float | None:
        """Parse Land Value / Acres from string (e.g. '273,300' or '0.38') to float."""
        if value is None:
            return None
        try:
            s = str(value).strip().replace(",", "")
            return float(s) if s else None
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _normalize_parcel_id(value: Any) -> str:
        """Normalize for comparison: digits only (e.g. 2463162679)."""
        s = str(value).strip()
        return re.sub(r"\D", "", s)

    @staticmethod
    def _extract_parcel_id_from_question(question: str) -> str | None:
        """Extract a parcel-ID-like token (e.g. 246316-2679 or 2463162679)."""
        # Match patterns like 246316-2679 or 2463162679 (6+ digits, optional hyphen, 4+ digits)
        m = re.search(r"\b(\d{6,}-?\d{4,})\b", question.strip())
        if m:
            return m.group(1).strip()
        return None

    def query_by_parcel_id(self, parcel_id: str) -> List[RAGChunk]:
        """Return chunks whose Parcel ID matches (exact or normalized)."""
        if not self._chunks:
            return []
        want = self._normalize_parcel_id(parcel_id)
        if not want:
            return []
        out: List[RAGChunk] = []
        for c in self._chunks:
            raw = c.metadata.get("Parcel ID") or c.metadata.get("Parcel ID ")
            if raw is None:
                continue
            if want == self._normalize_parcel_id(raw):
                out.append(c)
        return out

    @staticmethod
    def _parse_land_value_filter(question: str) -> Tuple[str | None, float | None, float | None]:
        """
        Parse question for land value comparison. Returns (operator, value1, value2).
        Operators: '>=', '<=', '>', '<', '==', 'between'
        - Single: (">=", 200000, None), ("<=", 100000, None), ("==", 200000, None)
        - Range: ("between", 100000, 200000)
        Supports: under/below X, equal to X, above/at least X, value range X to Y.
        """
        q = question.lower().strip()
        num_pattern = r"[\d,]+(?:\.[\d]+)?"

        def num(s: str) -> float:
            return float(s.replace(",", ""))

        # Value range: "value range 100000 to 200000", "between 100000 and 200000", "land value between X and Y"
        m = re.search(
            r"(?:land\s+value|value|price)\s+(?:range\s+)?(" + num_pattern + r")\s*(?:to|-|and)\s*(" + num_pattern + r")(?:\s*\$)?",
            q,
            re.I,
        )
        if not m:
            m = re.search(
                r"(?:between|from)\s+\$?\s*(" + num_pattern + r")\s+(?:and|to|-)\s+\$?\s*(" + num_pattern + r")(?:\s*\$)?",
                q,
                re.I,
            )
        if m:
            try:
                low, high = num(m.group(1)), num(m.group(2))
                if low > high:
                    low, high = high, low
                return ("between", low, high)
            except ValueError:
                pass

        # Equal to: "equal to 200000", "value equal to 200000", "exactly 200000", "land value exactly 200000"
        m = re.search(
            r"(?:equal\s+to|exactly)\s+\$?\s*(" + num_pattern + r")(?:\s*\$)?",
            q,
            re.I,
        )
        if not m:
            m = re.search(
                r"(?:land\s+value|value)\s+(?:equal\s+to|exactly)\s+(" + num_pattern + r")",
                q,
                re.I,
            )
        if m:
            try:
                return ("==", num(m.group(1)), None)
            except ValueError:
                pass

        # Under / below: "under 200000", "below 200000", "value under 200000$", "price below X"
        m = re.search(
            r"(?:below|under|less(?:\s+than)?|at\s+most|maximum)\s+(?:\w+\s+)*\$?\s*(" + num_pattern + r")(?:\s*\$)?",
            q,
            re.I,
        )
        if m:
            try:
                return ("<=", num(m.group(1)), None)
            except ValueError:
                pass
        m = re.search(r"<\s*\$?\s*(" + num_pattern + r")", q)
        if m:
            try:
                return ("<", num(m.group(1)), None)
            except ValueError:
                pass

        # Above / at least: "above 200000", "at least 200000", "value above 200000$"
        m = re.search(
            r"(?:above|over|greater(?:\s+than)?|at\s+least|minimum|more\s+than)\s+(?:\w+\s+)*\$?\s*(" + num_pattern + r")(?:\s*\$)?",
            q,
            re.I,
        )
        if not m:
            m = re.search(r"(" + num_pattern + r")\s*\$?\s*(?:and\s+)?(?:above|or\s+more)", q, re.I)
        if m:
            try:
                return (">=", num(m.group(1)), None)
            except ValueError:
                pass
        m = re.search(r">\s*\$?\s*(" + num_pattern + r")", q)
        if m:
            try:
                return (">", num(m.group(1)), None)
            except ValueError:
                pass

        # Plain "land value 200000" / "value 200000$" → treat as at least (>=)
        m = re.search(
            r"(?:land\s+value|value)\s*\$?\s*(" + num_pattern + r")(?:\s*\$)?",
            q,
            re.I,
        )
        if m:
            try:
                return (">=", num(m.group(1)), None)
            except ValueError:
                pass

        return (None, None, None)

    @staticmethod
    def _parse_land_use_filter(question: str) -> str | None:
        """Parse question for land use / use code (e.g. vacant, educational). Returns keyword or None."""
        q = question.lower().strip()
        keywords = [
            "vacant",
            "educational",
            "residential",
            "commercial",
            "industrial",
            "agricultural",
            "recreational",
            "forest",
        ]
        for kw in keywords:
            if kw in q:
                return kw
        return None

    @staticmethod
    def _parse_acre_filter(question: str) -> Tuple[str | None, float | None, float | None]:
        """
        Parse question for acre range. Returns (operator, value1, value2).
        - Single: (">=", 0.5, None), ("<=", 1.0, None), (">", 0.25, None), ("<", 2.0, None)
        - Range: ("between", 0.3, 1.0)
        Supports: acres above 0.5, over 1 acre, below 2 acres, between 0.3 and 1 acre, acre range 0.5 to 2.
        """
        q = question.lower().strip()
        num_pattern = r"[\d.]+"
        # Between / range: "between 0.3 and 1 acre", "0.5 to 2 acres", "acre range 0.3 - 1", "b/w 1.5 to 3 acres"
        m = re.search(
            r"(?:between|from)\s+(" + num_pattern + r")\s+(?:and|to|-)\s+(" + num_pattern + r")\s*(?:acre|acres)?",
            q,
            re.I,
        )
        if not m:
            m = re.search(
                r"(?:b/w|bw|btw)\s+(" + num_pattern + r")\s*(?:and|to|-)\s*(" + num_pattern + r")\s*(?:acre|acres)?",
                q,
                re.I,
            )
        if not m:
            m = re.search(
                r"(?:acre\s+)?range\s+(" + num_pattern + r")\s*(?:to|-)\s+(" + num_pattern + r")",
                q,
                re.I,
            )
        if m:
            try:
                low, high = float(m.group(1)), float(m.group(2))
                if low > high:
                    low, high = high, low
                return ("between", low, high)
            except ValueError:
                pass
        # Above / over / at least / minimum acres
        m = re.search(
            r"(?:acre(?:s)?\s+)?(?:above|over|greater(?:\s+than)?|at\s+least|minimum|more\s+than)\s+(" + num_pattern + r")\s*(?:acre|acres)?",
            q,
            re.I,
        )
        if not m:
            m = re.search(r"(" + num_pattern + r")\s*\+\s*(?:acre|acres)", q, re.I)
        if m:
            try:
                return (">=", float(m.group(1)), None)
            except ValueError:
                pass
        # Below / under / less than / maximum acres
        m = re.search(
            r"(?:acre(?:s)?\s+)?(?:below|under|less(?:\s+than)?|at\s+most|maximum)\s+(" + num_pattern + r")\s*(?:acre|acres)?",
            q,
            re.I,
        )
        if m:
            try:
                return ("<=", float(m.group(1)), None)
            except ValueError:
                pass
        m = re.search(r"acre(?:s)?\s*>\s*(" + num_pattern + r")", q, re.I)
        if m:
            try:
                return (">", float(m.group(1)), None)
            except ValueError:
                pass
        m = re.search(r"acre(?:s)?\s*<\s*(" + num_pattern + r")", q, re.I)
        if m:
            try:
                return ("<", float(m.group(1)), None)
            except ValueError:
                pass
        return (None, None, None)

    @staticmethod
    def _parse_requested_amount(question: str) -> int | None:
        """
        Parse requested count from user text, including formats like:
        - "50 rows", "give me 30"
        - "2k records", "1.5k rows"
        """
        q = (question or "").lower().strip()
        if not q:
            return None

        # e.g. 2k, 1.5k
        m = re.search(r"\b(\d+(?:\.\d+)?)\s*k\s*(?:data|records?|rows?|details?|entries)?\b", q, re.I)
        if m:
            try:
                return max(1, min(int(float(m.group(1)) * 1000), 2000))
            except ValueError:
                pass

        # e.g. 50 rows / 50 records
        m = re.search(r"\b(\d+)\s*(?:data|records?|rows?|details?|entries)\b", q, re.I)
        if m:
            try:
                return max(1, min(int(m.group(1)), 2000))
            except ValueError:
                pass

        # e.g. give me 50
        m = re.search(r"(?:give me|show me|get me|list|fetch)\s*(\d+)", q, re.I)
        if m:
            try:
                return max(1, min(int(m.group(1)), 2000))
            except ValueError:
                pass

        return None

    @staticmethod
    def _to_csv_text(rows: List[Dict[str, Any]], max_rows: int) -> str:
        """Render rows as CSV text for export-like requests."""
        use_rows = rows[:max_rows]
        if not use_rows:
            return "No rows available."

        headers: List[str] = []
        seen = set()
        for r in use_rows:
            for k in r.keys():
                ks = str(k)
                if ks not in seen:
                    seen.add(ks)
                    headers.append(ks)

        buf = io.StringIO()
        writer = csv.DictWriter(buf, fieldnames=headers, extrasaction="ignore")
        writer.writeheader()
        for r in use_rows:
            writer.writerow({h: r.get(h, "") for h in headers})
        return buf.getvalue().strip()

    @staticmethod
    def _rows_to_brief_lines(rows: List[Dict[str, Any]], max_rows: int) -> str:
        """Render deterministic row list for 'give me N rows' style asks."""
        use_rows = rows[:max_rows]
        if not use_rows:
            return "No matching rows found."
        lines: List[str] = []
        for idx, r in enumerate(use_rows, start=1):
            parcel = r.get("Parcel ID") or r.get("id") or r.get("ID") or "-"
            owner = r.get("Owner Placeholder Name") or r.get("owner_name") or "-"
            acres = r.get("Acres") or "-"
            lv = r.get("Land Value") or r.get("value") or "-"
            addr = r.get("Address") or "-"
            contact = r.get("Contact Info") or r.get("email") or "-"
            lines.append(
                f"{idx}. Parcel ID: {parcel} | Owner: {owner} | Acres: {acres} | Land Value: {lv} | Address: {addr} | Contact: {contact}"
            )
        return "\n".join(lines)

    @staticmethod
    def _parse_owner_filter(question: str) -> str | None:
        """
        Parse question for owner name. Returns the owner name string or None.
        Supports: owner X, owned by X, owner name X, parcels for X, for owner X, owner: X.
        """
        q = question.strip()
        if not q:
            return None
        # Quoted name: "owner 'John Smith'" or "owner \"Owner_0001\""
        m = re.search(r"(?:owner|owned\s+by|for\s+owner)\s*:?\s*[\'\"]([^\'\"]+)[\'\"]", q, re.I)
        if m:
            name = m.group(1).strip()
            if len(name) > 2:
                return name
        # After "owner", "owned by", "details for", "info for", "parcels for", "for owner" - capture name until stop words or end
        m = re.search(
            r"(?:owner(?:\s+name)?|owned\s+by|details?\s+for|info\s+for|parcels?\s+for|for\s+owner)\s*:?\s+([^\n]+?)(?=\s+with\s+|\s+above\s+|\s+below\s+|\s+and\s+\d|\s+over\s+\d|$)",
            q,
            re.I,
        )
        if m:
            name = m.group(1).strip()
            name = re.sub(r"\s+", " ", name)
            if len(name) > 0 and len(name) < 120:
                return name
        # Single word after "owner" / "owned by" (e.g. "owner Owner_0001")
        m = re.search(r"(?:owner|owned\s+by)\s+(\S+)", q, re.I)
        if m:
            return m.group(1).strip()
        return None

    def query_by_comparison(
        self,
        question: str,
        land_value_op: str | None,
        land_value_num: float | None,
        land_value_high: float | None,
        land_use_keyword: str | None,
        acre_op: str | None = None,
        acre_val: float | None = None,
        acre_high: float | None = None,
        owner_name: str | None = None,
        max_results: int = 25,
    ) -> List[RAGChunk]:
        """Filter chunks by land value, land use, acre range, and/or owner name. Returns matching rows for comparison."""
        if not self._chunks:
            return []
        out: List[RAGChunk] = []
        land_use_col = "Land Use / Use Code"
        land_value_col = "Land Value"
        acres_col = "Acres"
        owner_col = "Owner Placeholder Name"
        for c in self._chunks:
            meta = c.metadata
            # Owner filter: match Owner Placeholder Name and/or junction "owners" list
            if owner_name:
                raw_owner = meta.get(owner_col) or ""
                if owner_name.lower() in str(raw_owner).lower():
                    pass  # match
                else:
                    owners_list = meta.get("owners")
                    if isinstance(owners_list, list):
                        if not any(
                            owner_name.lower() in str(o.get("owner_name") or "").lower()
                            for o in owners_list
                            if isinstance(o, dict)
                        ):
                            continue
                    else:
                        continue
            # Land value filter (values in JSON may be strings like "273,300")
            if land_value_op is not None and land_value_num is not None:
                val = self._to_float(meta.get(land_value_col))
                if val is None:
                    continue
                if land_value_op == ">=" and val < land_value_num:
                    continue
                if land_value_op == ">" and val <= land_value_num:
                    continue
                if land_value_op == "<=" and val > land_value_num:
                    continue
                if land_value_op == "<" and val >= land_value_num:
                    continue
                if land_value_op == "==" and val != land_value_num:
                    continue
                if land_value_op == "between" and land_value_high is not None:
                    if not (land_value_num <= val <= land_value_high):
                        continue
            # Land use filter
            if land_use_keyword:
                raw_use = meta.get(land_use_col) or ""
                if land_use_keyword not in str(raw_use).lower():
                    continue
            # Acre filter
            if acre_op is not None and acre_val is not None:
                ac = self._to_float(meta.get(acres_col))
                if ac is None:
                    continue
                if acre_op == "between" and acre_high is not None:
                    if not (acre_val <= ac <= acre_high):
                        continue
                elif acre_op == ">=" and ac < acre_val:
                    continue
                elif acre_op == ">" and ac <= acre_val:
                    continue
                elif acre_op == "<=" and ac > acre_val:
                    continue
                elif acre_op == "<" and ac >= acre_val:
                    continue
            out.append(c)
        # Sort by Acres descending, then Land Value descending
        def sort_key(chunk: RAGChunk) -> Tuple[float, float]:
            a_f = self._to_float(chunk.metadata.get(acres_col)) or 0.0
            v_f = self._to_float(chunk.metadata.get(land_value_col)) or 0.0
            return (a_f, v_f)

        out.sort(key=sort_key, reverse=True)
        return out[:max_results]

    def query(self, question: str, top_k: int = 5) -> List[RAGChunk]:
        """
        Retrieve the top_k most similar rows. Tries in order:
        1) Parcel ID exact match
        2) Land value / land use comparison (above 200000, vacant land, etc.)
        3) Semantic search
        """
        if not self._chunks:
            raise RuntimeError("RAG index is empty. Did you call build_index()?")

        pid = self._extract_parcel_id_from_question(question)
        if pid:
            by_id = self.query_by_parcel_id(pid)
            if by_id:
                return by_id

        land_value_op, land_value_num, land_value_high = self._parse_land_value_filter(question)
        land_use_kw = self._parse_land_use_filter(question)
        acre_op, acre_val, acre_high = self._parse_acre_filter(question)
        owner_name = self._parse_owner_filter(question)
        if land_value_op is not None or land_use_kw or acre_op is not None or owner_name:
            by_comp = self.query_by_comparison(
                question,
                land_value_op=land_value_op,
                land_value_num=land_value_num,
                land_value_high=land_value_high,
                land_use_keyword=land_use_kw,
                acre_op=acre_op,
                acre_val=acre_val,
                acre_high=acre_high,
                owner_name=owner_name,
                max_results=30,
            )
            if by_comp:
                return by_comp[:20]

        q_emb = self._embed(question)
        q_vec = np.array(q_emb, dtype="float32")
        q_norm = np.linalg.norm(q_vec) or 1.0

        scored: List[Tuple[float, RAGChunk]] = []
        for chunk in self._chunks:
            c_vec = np.array(chunk.embedding, dtype="float32")
            c_norm = np.linalg.norm(c_vec) or 1.0
            score = float(np.dot(q_vec, c_vec) / (q_norm * c_norm))
            if not math.isnan(score):
                scored.append((score, chunk))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [c for _, c in scored[:top_k]]

    @staticmethod
    def _looks_weak_answer(answer: str) -> bool:
        """
        Detect generic/non-grounded answers and trigger a retry with broader context.
        """
        a = (answer or "").strip().lower()
        if not a:
            return True
        weak_markers = (
            "i don't have access",
            "i do not have access",
            "i'm not able to access",
            "cannot access the file",
            "could you provide more context",
            "could you share",
            "i couldn't find matching data",
            "i can't find matching data",
            "not enough information",
            "i am not sure",
            "i'm not sure",
            "unfortunately",
        )
        return any(m in a for m in weak_markers)

    @staticmethod
    def _merge_unique_chunks(primary: List["RAGChunk"], secondary: List["RAGChunk"], max_total: int = 12) -> List["RAGChunk"]:
        """
        Merge two chunk lists by unique row_index, preserving order preference.
        """
        out: List[RAGChunk] = []
        seen: set[int] = set()
        for chunk in primary + secondary:
            if chunk.row_index in seen:
                continue
            out.append(chunk)
            seen.add(chunk.row_index)
            if len(out) >= max_total:
                break
        return out

    @staticmethod
    def _is_behavioral_or_smalltalk_question(question: str) -> bool:
        """Detect casual chat that should not be forced through dataset retrieval."""
        q = (question or "").strip().lower()
        if not q:
            return False

        q_words = set(re.findall(r"[a-z0-9]+", q))

        def has_marker(marker: str) -> bool:
            marker = marker.lower().strip()
            if not marker:
                return False
            # Phrase marker: use substring match.
            if " " in marker:
                return marker in q
            # Single word marker: require whole-word match (avoid "hi" in "this").
            return marker in q_words

        # If data intent is present, keep the normal RAG path.
        data_markers = (
            "record", "records", "row", "rows", "dataset", "table", "data", "count",
            "how many", "list", "show", "filter", "owner", "parcel", "acre",
            "land value", "zoning", "json", "csv", "column", "file", "pdf", "docx",
            "document", "upload", "attached", "attachment",
        )
        if any(has_marker(marker) for marker in data_markers):
            return False

        chat_markers = (
            "hello", "hi", "hey", "how are you", "what are you doing", "who are you",
            "what can you do", "thank you", "thanks", "good morning", "good evening",
            "are you there", "what's up", "how's it going", "do you remember me",
        )
        return any(has_marker(marker) for marker in chat_markers)

    async def answer(
        self,
        question: str,
        top_k: int = 5,
        conversation_history: List[Dict[str, str]] | None = None,
        user_name: str | None = None,
    ) -> Dict[str, Any]:
        """
        High-level helper: retrieve relevant rows and have the LLM answer
        based on them. Optionally include prior conversation so the model
        understands question type and what was already asked.
        """
        if self._is_behavioral_or_smalltalk_question(question):
            history_prefix = ""
            if conversation_history:
                lines = []
                for turn in conversation_history[-MAX_HISTORY_TURNS:]:
                    role = (turn.get("role") or "").lower()
                    content = (turn.get("content") or "").strip()
                    if not content:
                        continue
                    label = "User" if role == "user" else "Assistant"
                    lines.append(f"{label}: {content[:MAX_HISTORY_CHARS]}")
                if lines:
                    history_prefix = "Conversation so far:\n" + "\n".join(lines) + "\n\n"

            system_prompt = BEHAVIORAL_CHAT_SYSTEM_PROMPT
            personalisation = get_personalisation_system_snippet(user_name)
            if personalisation:
                system_prompt = system_prompt.rstrip() + "\n\n" + personalisation

            client = LLMClient()
            answer = await client.complete(
                model="llama-3.1-8b-instant",
                system_prompt=system_prompt,
                user_prompt=f"{history_prefix}Current user message: {question}",
                temperature=0.4,
                max_tokens=120,
            )
            return {"answer": answer, "contexts": []}

        q_lower = question.lower().strip()
        q_words = re.findall(r"[a-z0-9]+", q_lower)
        relation_question = any(
            phrase in q_lower
            for phrase in (
                "relation",
                "relationship",
                "relationships",
                "related",
                "join",
                "foreign key",
                "linked",
                "mapping",
                "how tables connect",
                "how are these connected",
                "data model",
                "type of relation",
                "how does the data relate",
                "structure of the data",
                "how is the data organized",
                "what relations",
                "technical term",
                "social term",
                "business meaning",
            )
        )
        # "What is this file about?" / "What's inside?" / "Describe the file" / "What is this project about?"
        overview_question = any(
            phrase in q_lower
            for phrase in (
                "what is this file",
                "what is the file",
                "what's in this file",
                "what is in this file",
                "what does this file contain",
                "what does the file contain",
                "describe the file",
                "describe this file",
                "what is this dataset",
                "what is this data",
                "what is this table",
                "what is this project about",
                "what is the project about",
                "what is this about",
                "what's inside",
                "what is inside",
                "what kind of data",
                "what type of data",
                "tell me about this file",
                "explain this file",
                "summary of the file",
                "file summary",
                "overview of the file",
                "what information",
                "what does it contain",
            )
        )
        wants_all_details = any(
            phrase in q_lower
            for phrase in (
                "all detail", "all data", "every record", "full detail", "list all",
                "give me all", "show all", "all record", "all the data", "full data",
                "details of all", "all rows", "entire file", "whole file",
            )
        )
        property_browse_question = (
            "property" in q_lower
            and any(
                phrase in q_lower
                for phrase in (
                    "some property",
                    "about property",
                    "about some property",
                    "tell me about",
                    "show property",
                    "give property",
                )
            )
        )
        long_answer_intent = any(
            phrase in q_lower
            for phrase in (
                "explain in detail",
                "detailed explanation",
                "step by step",
                "why does",
                "why is",
                "analysis",
                "deep dive",
                "elaborate",
                "comprehensive",
            )
        )
        # Detect "give me N data", "2k records", etc.
        requested_n: int | None = None
        if not wants_all_details:
            requested_n = self._parse_requested_amount(q_lower)

        csv_request = any(k in q_lower for k in ("csv", "download csv", "export csv"))
        owner_list_request = any(k in q_lower for k in ("all owner", "land owner", "owner list", "all the owner"))
        generic_list_request = (
            any(k in q_lower for k in ("list", "show", "give me", "fetch"))
            and not csv_request
        )
        avg_value_request = (
            any(k in q_lower for k in ("avg", "average", "mean"))
            and any(k in q_lower for k in ("land value", "property", "cost", "rate", "price"))
        )
        if requested_n is not None:
            effective_k = min(requested_n, len(self._chunks), 100)
            effective_k = max(effective_k, top_k)
        elif wants_all_details:
            effective_k = min(50, len(self._chunks))
            effective_k = max(effective_k, top_k)
        elif overview_question:
            # For "what is this file about?" use a few sample rows so the model can describe content
            effective_k = min(10, len(self._chunks))
            effective_k = max(effective_k, top_k)
        else:
            # Slightly broader retrieval helps open-ended DB questions.
            effective_k = max(top_k, 8)
        # Stay under Groq token-per-request limit (e.g. 6k); cap rows sent in one call
        effective_k = min(effective_k, MAX_ROWS_PER_REQUEST)
        # For "what is this file about?" use first N rows so we always have a concrete sample to describe
        if overview_question and self._chunks:
            chunks = sorted(self._chunks, key=lambda c: c.row_index)[:effective_k]
        else:
            chunks = self.query(question=question, top_k=effective_k)

        if not chunks:
            return {
                "answer": "I couldn't find matching data for your question.",
                "contexts": [],
            }

        # Reuse filter parsing for deterministic utilities (csv/list/avg).
        land_value_op, land_value_num, land_value_high = self._parse_land_value_filter(question)
        land_use_kw = self._parse_land_use_filter(question)
        acre_op, acre_val, acre_high = self._parse_acre_filter(question)
        owner_name = self._parse_owner_filter(question)

        # Deterministic response paths for DB-style utility asks.
        comparison_rows = self.query_by_comparison(
            question,
            land_value_op=land_value_op,
            land_value_num=land_value_num,
            land_value_high=land_value_high,
            land_use_keyword=land_use_kw,
            acre_op=acre_op,
            acre_val=acre_val,
            acre_high=acre_high,
            owner_name=owner_name,
            max_results=max(2000, requested_n or 50),
        )
        selected_rows = [c.metadata for c in (comparison_rows if comparison_rows else self._chunks)]

        if avg_value_request:
            vals = [self._to_float(r.get("Land Value")) for r in selected_rows]
            nums = [v for v in vals if v is not None]
            if nums:
                avg_v = sum(nums) / len(nums)
                min_v = min(nums)
                max_v = max(nums)
                return {
                    "answer": (
                        f"Average Land Value: {avg_v:,.2f} (based on {len(nums)} rows). "
                        f"Range: {min_v:,.0f} to {max_v:,.0f}."
                    ),
                    "contexts": [{"row_index": c.row_index, "metadata": c.metadata, "text": c.text} for c in chunks],
                }

        if csv_request:
            req = requested_n or len(selected_rows)
            csv_text = self._to_csv_text(selected_rows, max_rows=min(req, 2000))
            return {
                "answer": csv_text,
                "contexts": [{"row_index": c.row_index, "metadata": c.metadata, "text": c.text} for c in chunks],
            }

        range_query_request = (
            land_value_op is not None
            or land_use_kw is not None
            or acre_op is not None
            or owner_name is not None
        )

        if wants_all_details or requested_n is not None or owner_list_request or generic_list_request or range_query_request:
            # Priority:
            # 1) explicit requested amount (e.g. 25, 2k)
            # 2) for range/filter queries, return all matching rows in view
            # 3) otherwise, keep a practical default for generic list asks
            if requested_n is not None:
                req = requested_n
            elif range_query_request:
                req = len(selected_rows)
            else:
                req = 50 if generic_list_request else len(selected_rows)
            req = min(req, 2000)
            lines = self._rows_to_brief_lines(selected_rows, max_rows=req)
            return {
                "answer": f"Showing {min(req, len(selected_rows))} rows:\n{lines}",
                "contexts": [{"row_index": c.row_index, "metadata": c.metadata, "text": c.text} for c in chunks],
            }

        context_blocks = []
        for c in chunks:
            row_text = c.text
            # For document-like uploads, allow a larger snippet per chunk.
            row_text_cap = 1000 if ("content:" in row_text and len(row_text) > 450) else MAX_ROW_TEXT_CHARS
            if len(row_text) > row_text_cap:
                row_text = row_text[:row_text_cap].rstrip() + "..."
            context_blocks.append(f"- ROW {c.row_index}: {row_text}")

        context_str = "\n".join(context_blocks)
        total_rows = getattr(self, "_total_row_count", None) or len(self._chunks)
        schema_summary = getattr(self, "_schema_summary", "").strip()

        count_info = (
            f"Total data rows in the file/table: {total_rows}. "
            f"When the user asks 'how many records/rows', answer with this exact number: {total_rows}. Do not say 'Yes' or use the sample row count.\n\n"
        )

        wants_many_rows = wants_all_details or (requested_n is not None)
        is_comparison = len(chunks) > 5 and not wants_many_rows
        short_answer_intent = (
            not relation_question
            and not overview_question
            and not wants_many_rows
            and not property_browse_question
            and not long_answer_intent
            and len(q_words) <= 12
            and (
                q_lower.endswith("?")
                or any(
                    q_lower.startswith(prefix)
                    for prefix in (
                        "is ",
                        "are ",
                        "can ",
                        "do ",
                        "does ",
                        "did ",
                        "what ",
                        "who ",
                        "when ",
                        "where ",
                        "which ",
                        "how many ",
                    )
                )
            )
        )
        if relation_question:
            instruction = (
                "The user is asking about relationships or how the data is structured. "
                "Explain in two ways: "
                "(1) Technical: which columns look like keys or IDs (e.g. *._id, *.id), likely foreign-key links, one-to-many or id-to-name pairs, and how you would join or link tables. "
                "(2) Social/plain language: what the relationships mean in real-world or business terms (e.g. 'A Project has an Owner', 'Contact is the person on the lead', 'Account is the company or customer'). "
                "Use the schema summary and the row examples above. If a link is only inferred, say 'likely' or 'appears to'. Keep it clear and concise."
            )
        elif overview_question:
            instruction = (
                "The user is asking what this file or dataset is about. Using ONLY the schema summary and total row count above, and the sample rows below, describe: "
                "(1) what kind of data this is (e.g. projects, contacts, parcels, CRM records); "
                "(2) the main columns and what they represent; "
                "(3) how many records there are. "
                "You may briefly mention 1–2 example values from the sample rows. Be specific and base your answer only on the information provided—do not say you don't have access or don't know; you have the schema and sample rows."
            )
        elif property_browse_question:
            instruction = (
                "The user is asking generally about properties. Provide a short overview and then give 3 concrete property examples from the rows below. "
                "For each example include Parcel ID, Address, Land Value, Land Use, and Owner if present."
            )
        elif wants_all_details:
            instruction = (
                "The user asked for ALL or FULL details. List or summarize the rows below. Do not reply with only the record count. Say something like 'Here are details from your file' or 'Summary of the data' and then list the key columns/values for as many rows as shown below. You can group or summarize if there are many; include at least several concrete examples."
            )
        elif requested_n is not None:
            cap_note = f" (Showing up to {len(chunks)} rows per request due to API limits; user asked for {requested_n}.)" if requested_n > len(chunks) else ""
            instruction = (
                f"The user asked for {requested_n} data/records.{cap_note} List every row shown below (up to {len(chunks)} rows). Do not summarize or skip rows—output each row's details in order (e.g. ROW 1: ..., ROW 2: ..., etc.). Do not reply with only the record count."
            )
        elif is_comparison:
            instruction = (
                "This is a comparison/filter result. Summarize: total count of matching rows; key columns and ranges when present; and 2-3 example rows. Keep to a short paragraph or bullets."
            )
        elif short_answer_intent:
            instruction = (
                "This is a short question. Return only a short direct answer (one line, max ~20 words). "
                "No extra explanation unless the user asks for details."
            )
        else:
            instruction = (
                "Match your answer to what was asked. "
                "Only for 'how many records/rows': give the total row count above and do not say 'Yes'. "
                "For 'give me N details' or 'list': provide those details/rows, do not only state the count. "
                "For yes/no questions: reply with 'Yes.' or 'No.' first. Otherwise answer in 1-2 sentences. Be concise."
            )

        # Prepend conversation history so the model knows what was already asked
        conversation_prefix = ""
        if conversation_history:
            lines = []
            for turn in conversation_history[-MAX_HISTORY_TURNS:]:
                role = (turn.get("role") or "").lower()
                content = (turn.get("content") or "").strip()
                if not content:
                    continue
                if len(content) > MAX_HISTORY_CHARS:
                    content = content[:MAX_HISTORY_CHARS].rstrip() + "..."
                label = "User" if role == "user" else "Assistant"
                lines.append(f"{label}: {content}")
            if lines:
                conversation_prefix = "Conversation so far:\n" + "\n".join(lines) + "\n\n"

        rows_label = "Rows from the table" if wants_many_rows else "Sample rows (use total row count above only for 'how many' questions)"
        user_prompt = (
            f"{conversation_prefix}"
            f"Current question: {question}\n\n"
            f"{count_info}"
            f"Schema/relationship summary:\n{schema_summary or 'Schema summary unavailable.'}\n\n"
            f"{rows_label}:\n{context_str}\n\n"
            f"{instruction}"
        )

        base_system_prompt = TABLE_DATASET_SYSTEM_PROMPT if self._from_user_table else PARCEL_DATASET_SYSTEM_PROMPT
        system_prompt = base_system_prompt
        if conversation_history:
            system_prompt = system_prompt.rstrip() + "\n\n" + CONVERSATION_AND_QUESTION_TYPE_INSTRUCTION.strip()
        personalisation = get_personalisation_system_snippet(user_name)
        if personalisation:
            system_prompt = system_prompt.rstrip() + "\n\n" + personalisation
        try:
            client = LLMClient()
            # Allow long answers when user asks for relationships, full details, overview, or many rows
            if relation_question:
                max_tok = 2048  # technical + social explanation can be long
            elif overview_question:
                max_tok = 512  # file/dataset description: schema + examples
            elif short_answer_intent:
                max_tok = 96  # short question => short response
            elif wants_many_rows:
                max_tok = 2048  # listing many rows needs more output
            elif wants_all_details:
                max_tok = 2048  # "all details" / "full data" can be long
            elif is_comparison:
                max_tok = 400
            else:
                max_tok = 256
            answer = await client.complete(
                model="llama-3.1-8b-instant",
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                temperature=0.2,
                max_tokens=max_tok,
            )

            # If first pass is weak/generic, retry once with broader context and stricter grounding.
            if self._looks_weak_answer(answer) and self._chunks:
                try:
                    wider_k = min(max(len(chunks) + 4, 10), MAX_ROWS_PER_REQUEST)
                    broader_chunks = self.query(question=question, top_k=wider_k)
                    retry_chunks = self._merge_unique_chunks(chunks, broader_chunks, max_total=MAX_ROWS_PER_REQUEST)

                    retry_blocks = []
                    for c in retry_chunks:
                        row_text = c.text
                        row_text_cap = 1000 if ("content:" in row_text and len(row_text) > 450) else MAX_ROW_TEXT_CHARS
                        if len(row_text) > row_text_cap:
                            row_text = row_text[:row_text_cap].rstrip() + "..."
                        retry_blocks.append(f"- ROW {c.row_index}: {row_text}")

                    retry_prompt = (
                        f"{conversation_prefix}"
                        f"Current question: {question}\n\n"
                        f"{count_info}"
                        f"Schema/relationship summary:\n{schema_summary or 'Schema summary unavailable.'}\n\n"
                        f"Rows from the table:\n" + "\n".join(retry_blocks) + "\n\n"
                        "Answer strictly from these rows. Do not say you lack access to file/data. "
                        "If exact value is unavailable, state what is available and provide closest relevant details."
                    )
                    answer_retry = await client.complete(
                        model="llama-3.1-8b-instant",
                        system_prompt=system_prompt,
                        user_prompt=retry_prompt,
                        temperature=0.2,
                        max_tokens=max_tok,
                    )
                    if not self._looks_weak_answer(answer_retry):
                        answer = answer_retry
                        chunks = retry_chunks
                except Exception:
                    pass
        except Exception as e:
            err_msg = str(e).lower()
            # Request too large for Groq (e.g. 413 / 6k TPM limit)
            if "413" in str(e) or "request too large" in err_msg or ("tokens" in err_msg and "limit" in err_msg):
                # Retry once with a much smaller prompt (fewer rows + no conversation prefix).
                try:
                    retry_count = max(3, min(6, len(chunks) // 2))
                    retry_chunks = chunks[:retry_count]
                    retry_blocks = []
                    for c in retry_chunks:
                        row_text = c.text
                        if len(row_text) > MAX_ROW_TEXT_CHARS:
                            row_text = row_text[:MAX_ROW_TEXT_CHARS].rstrip() + "..."
                        retry_blocks.append(f"- ROW {c.row_index}: {row_text}")
                    retry_user_prompt = (
                        f"Current question: {question}\n\n"
                        f"{count_info}"
                        f"Rows from the table:\n" + "\n".join(retry_blocks) + "\n\n"
                        "List the row details shown above. Keep output concise and factual."
                    )
                    retry_system_prompt = base_system_prompt
                    if personalisation:
                        retry_system_prompt = retry_system_prompt.rstrip() + "\n\n" + personalisation
                    client = LLMClient()
                    answer = await client.complete(
                        model="llama-3.1-8b-instant",
                        system_prompt=retry_system_prompt,
                        user_prompt=retry_user_prompt,
                        temperature=0.2,
                        max_tokens=1024,
                    )
                except Exception:
                    answer = (
                        "That request would send too much data in one go (API limit). "
                        "Try asking for fewer rows, e.g. 'give me 6 data' or 'show 5 records'."
                    )
            elif "closed" in err_msg or "connection" in err_msg or "nodename" in err_msg or "servname" in err_msg or "connecterror" in err_msg:
                try:
                    client = LLMClient()
                    answer = await client.complete(
                        model="llama-3.1-8b-instant",
                        system_prompt=system_prompt,
                        user_prompt=user_prompt,
                        temperature=0.2,
                        max_tokens=2048 if (relation_question or wants_many_rows or wants_all_details) else (400 if is_comparison else 256),
                    )
                except Exception:
                    answer = "Sorry, the request could not be completed. Please try again in a moment."
            else:
                raise

        return {
            "answer": answer,
            "contexts": [
                {"row_index": c.row_index, "metadata": c.metadata, "text": c.text}
                for c in chunks
            ],
        }


_rag_instance: ExcelRAG | None = None


# Default dataset: app/us_parcel_dataset_2000.json (RAG uses this; supports "owners" junction per parcel).
DEFAULT_DATASET_FILENAME = "us_parcel_dataset_2000.json"


def get_rag() -> ExcelRAG:
    """
    Singleton-style accessor for the dataset RAG index.

    Uses DATASET_PATH or DATASET_JSON_PATH if set; otherwise uses app/us_parcel_dataset_2000.json.
    All RAG answers are based on this dataset (JSON or CSV).
    """
    global _rag_instance

    if _rag_instance is None:
        dataset_path = os.getenv("DATASET_PATH") or os.getenv("DATASET_JSON_PATH")
        if not dataset_path:
            _project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
            app_path = os.path.join(_project_root, "app", DEFAULT_DATASET_FILENAME)
            root_path = os.path.join(_project_root, DEFAULT_DATASET_FILENAME)
            dataset_path = app_path if os.path.exists(app_path) else root_path
        if not os.path.exists(dataset_path):
            raise RuntimeError(
                f"Dataset file not found at {dataset_path}. "
                f"Set DATASET_PATH or DATASET_JSON_PATH or place {DEFAULT_DATASET_FILENAME} in app/ or project root."
            )

        rag = ExcelRAG(dataset_path=dataset_path)
        rag.build_index()
        _rag_instance = rag

    return _rag_instance


def _pdf_to_records(content: bytes) -> List[Dict[str, Any]]:
    """Extract tables or full text from PDF into list of record dicts."""
    import io
    import pdfplumber
    records: List[Dict[str, Any]] = []
    text_chunk_records: List[Dict[str, Any]] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    if not table:
                        continue
                    headers = [str(h or "").strip() or f"col_{i}" for i, h in enumerate(table[0])]
                    for row in table[1:]:
                        vals = [str(c or "").strip() for c in row]
                        if len(vals) < len(headers):
                            vals += [""] * (len(headers) - len(vals))
                        records.append(dict(zip(headers, vals[: len(headers)])))
            t = (page.extract_text() or "").strip()
            if t:
                for chunk_idx, chunk in enumerate(_chunk_text(t), start=1):
                    text_chunk_records.append(
                        {
                            "content": chunk,
                            "source_type": "pdf_text",
                            "page": page_num,
                            "chunk": chunk_idx,
                        }
                    )
    if records:
        return records
    if text_chunk_records:
        return text_chunk_records
    return []


def _docx_to_records(content: bytes) -> List[Dict[str, Any]]:
    """Extract tables or paragraphs from DOCX into list of record dicts."""
    import io
    from docx import Document
    doc = Document(io.BytesIO(content))
    records: List[Dict[str, Any]] = []
    for table in doc.tables:
        rows = list(table.rows)
        if not rows:
            continue
        headers = [str(cell.text or "").strip() or f"col_{i}" for i, cell in enumerate(rows[0].cells)]
        for row in rows[1:]:
            vals = [str(cell.text or "").strip() for cell in row.cells]
            if len(vals) < len(headers):
                vals += [""] * (len(headers) - len(vals))
            records.append(dict(zip(headers, vals[: len(headers)])))
    if records:
        return records
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if parts:
        joined = "\n\n".join(parts)
        chunks = _chunk_text(joined)
        return [
            {
                "content": chunk,
                "source_type": "docx_text",
                "chunk": idx + 1,
            }
            for idx, chunk in enumerate(chunks)
        ]
    return []


def parse_uploaded_file(content: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Parse uploaded file content into list of record dicts.
    Supports: .xlsx/.xls (Excel), .csv, .json, .txt (CSV/TSV), .pdf, .docx.
    """
    name_lower = (filename or "").lower()
    if name_lower.endswith(".pdf"):
        return _pdf_to_records(content)
    if name_lower.endswith(".docx"):
        return _docx_to_records(content)
    if name_lower.endswith(".xlsx") or name_lower.endswith(".xls") or name_lower.endswith(".xlsm"):
        import pandas as pd
        import io
        df = pd.read_excel(io.BytesIO(content))
        df = df.fillna("")
        return [dict(zip(df.columns, row)) for row in df.to_numpy().tolist()]
    # Text-based: decode and parse (CSV, JSON, TXT)
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content.decode("utf-8-sig")
        except Exception:
            text = content.decode("latin-1")
    text = text.lstrip("\ufeff")
    records = parse_table_input(text)
    # If text isn't tabular, still make it queryable as document chunks.
    if not records:
        chunks = _chunk_text(text)
        if chunks:
            return [
                {
                    "content": chunk,
                    "source_type": "text",
                    "chunk": idx + 1,
                }
                for idx, chunk in enumerate(chunks)
            ]
    return records


def get_rag_from_table(table: Union[str, List[Any]]) -> ExcelRAG:
    """
    Build a one-off RAG index from a user-provided table (CSV string, JSON string,
    or list of dicts / list of rows). Use this to answer questions over that table only.
    """
    records = parse_table_input(table)
    if not records:
        raise ValueError("Table could not be parsed or is empty. Provide CSV (first row = headers), JSON array of objects, or list of dicts.")
    rag = ExcelRAG(records=records)
    rag.build_index()
    return rag

